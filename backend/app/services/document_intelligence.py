from __future__ import annotations

import re
from datetime import datetime
from functools import lru_cache
from pathlib import Path

from docx import Document

from app.core.config import get_settings
from app.schemas.analysis import InsightResponse, SummaryLength, SummaryMode, SummaryResponse
from app.services.cache_service import CacheService


STOPWORDS = {
    "the", "and", "for", "that", "with", "this", "from", "have", "will", "your", "into", "their", "there", "which",
    "about", "would", "could", "should", "these", "those", "while", "where", "when", "been", "being", "than", "then",
    "them", "they", "also", "over", "after", "before", "such", "only", "very", "more", "most", "some", "much",
}


class DocumentIntelligenceService:
    """Summaries, keywords, topics, sentiment, and classification for uploaded documents."""

    def __init__(self) -> None:
        self.settings = get_settings()
        self.cache = CacheService()

    def summarize(
        self,
        *,
        job_id: str,
        source_type: str,
        text: str,
        mode: SummaryMode,
        length: SummaryLength,
        language_hint: str,
    ) -> SummaryResponse:
        if not text.strip():
            raise ValueError("There is no extractable text available to summarize.")

        cache_key = self.cache.build_key(job_id, source_type, mode, length, language_hint, text[:5000])
        cached = self.cache.get("summaries", cache_key)
        if cached is not None:
            return SummaryResponse(**cached)

        if mode == "abstractive":
            result = self._abstractive_summary(job_id, source_type, text, length, language_hint)
        elif mode == "bullet":
            result = self._bullet_summary(job_id, source_type, text, length, language_hint)
        else:
            result = self._extractive_summary(job_id, source_type, text, length, language_hint)

        self.cache.set("summaries", cache_key, result.model_dump())
        return result

    def analyze(self, *, job_id: str, source_type: str, text: str) -> InsightResponse:
        if not text.strip():
            raise ValueError("There is no extractable text available to analyze.")

        cache_key = self.cache.build_key(job_id, source_type, text[:5000])
        cached = self.cache.get("insights", cache_key)
        if cached is not None:
            return InsightResponse(**cached)

        keywords = self._keywords(text, top_k=12)
        topics = keywords[:5]
        sentiment_label, sentiment_score, sentiment_model, sentiment_fallback = self._sentiment(text)
        classification_label, classification_score, class_model, class_fallback = self._classify(text)
        response = InsightResponse(
            job_id=job_id,
            source_type=source_type,
            keywords=keywords,
            topics=topics,
            sentiment_label=sentiment_label,
            sentiment_score=sentiment_score,
            classification_label=classification_label,
            classification_score=classification_score,
            generated_at=datetime.utcnow(),
            model_used=f"{sentiment_model}; {class_model}",
            used_fallback=sentiment_fallback or class_fallback,
        )
        self.cache.set("insights", cache_key, response.model_dump(mode="json"))
        return response

    def export_summary(self, *, job_id: str, summary: SummaryResponse, output_dir: Path, format_type: str) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        base_name = f"{job_id}_{summary.mode}_{summary.length}_summary"
        if format_type == "txt":
            export_path = output_dir / f"{base_name}.txt"
            export_path.write_text(summary.summary_text, encoding="utf-8")
            return export_path

        export_path = output_dir / f"{base_name}.docx"
        document = Document()
        document.add_heading("AlignPDF Summary", level=1)
        document.add_paragraph(summary.summary_text)
        if summary.bullets:
            document.add_heading("Key Points", level=2)
            for bullet in summary.bullets:
                document.add_paragraph(bullet, style="List Bullet")
        document.save(export_path)
        return export_path

    def _extractive_summary(self, job_id: str, source_type: str, text: str, length: SummaryLength, language_hint: str) -> SummaryResponse:
        sentences = self._sentences(text)
        sentence_scores = self._rank_sentences(sentences)
        selected = self._select_ranked_sentences(sentence_scores, length)
        summary_text = " ".join(selected).strip()
        return SummaryResponse(
            job_id=job_id,
            source_type=source_type,
            mode="extractive",
            length=length,
            summary_text=summary_text,
            bullets=selected,
            language=language_hint,
            model_used="heuristic-extractive-ranker",
            used_fallback=False,
        )

    def _bullet_summary(self, job_id: str, source_type: str, text: str, length: SummaryLength, language_hint: str) -> SummaryResponse:
        sentences = self._sentences(text)
        ranked = self._select_ranked_sentences(self._rank_sentences(sentences), length)
        bullets = [sentence.strip() for sentence in ranked]
        return SummaryResponse(
            job_id=job_id,
            source_type=source_type,
            mode="bullet",
            length=length,
            summary_text="\n".join(f"- {sentence}" for sentence in bullets),
            bullets=bullets,
            language=language_hint,
            model_used="heuristic-bullet-ranker",
            used_fallback=False,
        )

    def _abstractive_summary(self, job_id: str, source_type: str, text: str, length: SummaryLength, language_hint: str) -> SummaryResponse:
        if self.settings.enable_transformers:
            try:
                summarizer = self._load_pipeline("summarization", self.settings.summary_default_model)
                token_limits = {"short": (50, 100), "medium": (100, 180), "long": (180, 260)}
                min_length, max_length = token_limits[length]
                generated = summarizer(text[:6000], min_length=min_length, max_length=max_length, truncation=True)[0]["summary_text"]
                bullets = [sentence.strip() for sentence in self._sentences(generated)[:6]]
                return SummaryResponse(
                    job_id=job_id,
                    source_type=source_type,
                    mode="abstractive",
                    length=length,
                    summary_text=generated.strip(),
                    bullets=bullets,
                    language=language_hint,
                    model_used=self.settings.summary_default_model,
                    used_fallback=False,
                )
            except Exception:
                pass

        fallback = self._extractive_summary(job_id, source_type, text, length, language_hint)
        fallback.mode = "abstractive"
        fallback.model_used = "heuristic-fallback-summary"
        fallback.used_fallback = True
        return fallback

    @staticmethod
    def _sentences(text: str) -> list[str]:
        return [sentence.strip() for sentence in re.split(r"(?<=[.!?])\s+", text.replace("\n", " ")) if sentence.strip()]

    def _rank_sentences(self, sentences: list[str]) -> list[tuple[float, str]]:
        frequencies = {}
        for word in self._words(" ".join(sentences)):
            frequencies[word] = frequencies.get(word, 0) + 1

        ranked = []
        for sentence in sentences:
            words = self._words(sentence)
            if not words:
                continue
            score = sum(frequencies.get(word, 0) for word in words) / len(words)
            ranked.append((score, sentence))
        return sorted(ranked, key=lambda item: item[0], reverse=True)

    @staticmethod
    def _select_ranked_sentences(ranked: list[tuple[float, str]], length: SummaryLength) -> list[str]:
        target = {"short": 3, "medium": 5, "long": 8}[length]
        selected = [sentence for _, sentence in ranked[:target]]
        return selected

    def _keywords(self, text: str, top_k: int) -> list[str]:
        frequencies = {}
        for word in self._words(text):
            frequencies[word] = frequencies.get(word, 0) + 1
        ranked = sorted(frequencies.items(), key=lambda item: (-item[1], item[0]))
        return [word for word, _ in ranked[:top_k]]

    @staticmethod
    def _words(text: str) -> list[str]:
        return [
            word
            for word in re.findall(r"[A-Za-z][A-Za-z\-]{2,}", text.lower())
            if word not in STOPWORDS
        ]

    def _sentiment(self, text: str) -> tuple[str, float, str, bool]:
        if self.settings.enable_transformers:
            try:
                classifier = self._load_pipeline("sentiment-analysis", self.settings.sentiment_default_model)
                result = classifier(text[:3000])[0]
                return result["label"].lower(), float(result["score"]), self.settings.sentiment_default_model, False
            except Exception:
                pass

        positive = sum(token in text.lower() for token in ("good", "great", "success", "improve", "benefit"))
        negative = sum(token in text.lower() for token in ("error", "fail", "risk", "issue", "problem"))
        if positive >= negative:
            return "positive", 0.55, "heuristic-sentiment", True
        return "negative", 0.55, "heuristic-sentiment", True

    def _classify(self, text: str) -> tuple[str, float, str, bool]:
        labels = ["report", "invoice", "academic", "legal", "manual", "form", "brochure", "presentation"]
        if self.settings.enable_transformers:
            try:
                classifier = self._load_pipeline("zero-shot-classification", self.settings.classification_default_model)
                result = classifier(text[:3000], labels)
                return result["labels"][0], float(result["scores"][0]), self.settings.classification_default_model, False
            except Exception:
                pass

        lowered = text.lower()
        if "abstract" in lowered or "references" in lowered:
            return "academic", 0.62, "heuristic-classifier", True
        if "invoice" in lowered or "amount due" in lowered:
            return "invoice", 0.68, "heuristic-classifier", True
        if "agreement" in lowered or "party" in lowered:
            return "legal", 0.6, "heuristic-classifier", True
        return "report", 0.51, "heuristic-classifier", True

    @staticmethod
    @lru_cache(maxsize=6)
    def _load_pipeline(task: str, model: str):
        from transformers import pipeline  # type: ignore

        return pipeline(task, model=model)
