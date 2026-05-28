from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document

from app.services.ocr_service import OcrService
from app.services.pdf_detector import PdfDetector


class TextExtractorService:
    """Extract text from PDFs and DOCX files for downstream AI analysis."""

    def __init__(self) -> None:
        self.detector = PdfDetector()
        self.ocr_service = OcrService()

    def extract_pdf_text(self, pdf_path: Path) -> str:
        detection = self.detector.inspect_pdf(pdf_path)
        text_pages: list[str] = []
        temp_dir = pdf_path.parent
        with fitz.open(pdf_path) as document:
            for index, page in enumerate(document):
                page_stats = detection["page_stats"][index]
                page_text = page.get_text("text").strip()
                if page_text:
                    text_pages.append(page_text)
                    continue
                ocr_layout = self.ocr_service.extract_page(page, page_stats, temp_dir)
                text_pages.append("\n".join(" ".join(span.text for span in block.spans) for block in ocr_layout.blocks))
        return "\n\n".join(page for page in text_pages if page.strip())

    @staticmethod
    def extract_docx_text(docx_path: Path) -> str:
        document = Document(docx_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        return "\n".join(parts)
