from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt, RGBColor

from app.schemas.jobs import ImageBlock, ParagraphBlock, PdfAnalysis, TableBlock
from app.utils.docx_helpers import ALIGNMENT_MAP, set_section_columns


class DocxBuilder:
    """Convert the normalized layout model into a best-effort DOCX replica."""

    def build(self, analysis: PdfAnalysis, output_path: Path) -> None:
        document = Document()
        self._set_default_style(document)

        for page_index, page in enumerate(analysis.pages):
            section = document.sections[-1] if page_index == 0 else document.add_section(WD_SECTION_START.NEW_PAGE)
            section.page_width = Pt(page.width)
            section.page_height = Pt(page.height)
            section.top_margin = Pt(page.margins["top"])
            section.bottom_margin = Pt(page.margins["bottom"])
            section.left_margin = Pt(page.margins["left"])
            section.right_margin = Pt(page.margins["right"])
            set_section_columns(section, page.columns)

            if page.header_text:
                header = section.header.paragraphs[0]
                header.text = page.header_text
            if page.footer_text:
                footer = section.footer.paragraphs[0]
                footer.text = page.footer_text

            for block in page.blocks:
                if isinstance(block, ParagraphBlock):
                    self._add_paragraph(document, block)
                elif isinstance(block, TableBlock):
                    self._add_table(document, block)
                elif isinstance(block, ImageBlock):
                    self._add_image(document, block)

        document.save(output_path)

    @staticmethod
    def _set_default_style(document: Document) -> None:
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

    def _add_paragraph(self, document: Document, block: ParagraphBlock) -> None:
        style_name = None
        if block.kind == "list":
            style_name = "List Bullet" if block.list_type == "bullet" else "List Number"
        paragraph = document.add_paragraph(style=style_name)
        paragraph.alignment = ALIGNMENT_MAP.get(block.alignment, ALIGNMENT_MAP["left"])
        paragraph.paragraph_format.left_indent = Pt(block.left_indent * 0.18)
        paragraph.paragraph_format.first_line_indent = Pt(block.first_line_indent * 0.18)
        paragraph.paragraph_format.space_before = Pt(block.space_before)
        paragraph.paragraph_format.space_after = Pt(block.space_after)
        paragraph.paragraph_format.line_spacing = block.line_spacing

        for span in block.spans:
            run = paragraph.add_run(span.text)
            run.font.name = span.style.font_name or "Calibri"
            run.font.size = Pt(span.style.font_size)
            run.bold = span.style.bold
            run.italic = span.style.italic
            run.underline = span.style.underline
            if span.style.color:
                try:
                    run.font.color.rgb = RGBColor.from_string(span.style.color[-6:].upper())
                except ValueError:
                    pass

    @staticmethod
    def _add_table(document: Document, block: TableBlock) -> None:
        row_count = len(block.rows)
        col_count = max((sum(cell.col_span for cell in row) for row in block.rows), default=1)
        table = document.add_table(rows=row_count, cols=col_count)
        table.style = "Table Grid"

        for row_index, row in enumerate(block.rows):
            col_index = 0
            for cell in row:
                target = table.cell(row_index, col_index)
                target.text = cell.text
                if cell.bold:
                    target.paragraphs[0].runs[0].bold = True
                if cell.italic:
                    target.paragraphs[0].runs[0].italic = True
                if cell.col_span > 1:
                    merge_target = table.cell(row_index, col_index + cell.col_span - 1)
                    target.merge(merge_target)
                if cell.row_span > 1 and row_index + cell.row_span - 1 < row_count:
                    merge_target = table.cell(row_index + cell.row_span - 1, col_index)
                    target.merge(merge_target)
                col_index += cell.col_span

    @staticmethod
    def _add_image(document: Document, block: ImageBlock) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(block.image_path, width=Pt(block.bbox[2] - block.bbox[0]), height=Pt(block.bbox[3] - block.bbox[1]))

